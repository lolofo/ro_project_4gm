import docx


def creat_doc_data(N,a,d,q) :
    '''
    creation a docx doc --> add just the data to the problem
    '''
    doc = docx.Document() 
  
    doc.add_heading('One machine sequencing problem', 0)
    doc.add_paragraph('Presentation of the data : ')

    data = []

    for i in range(len(N)) :
        job = "Job "+str(N[i])
        row = (job , a[i] , d[i] , q[i])
        data.append(row)
    
    data = tuple(data)

    table = doc.add_table(rows=1, cols=4)
    row = table.rows[0].cells 
    row[0].text = 'Job'
    row[1].text = 'starting_time'
    row[2].text = 'processing_time'
    row[3].text = 'queuing_time'

    for n,start , process,queuing in data: 
  
    
        row = table.add_row().cells 
        row[0].text = n 
        row[1].text = str(start)
        row[2].text = str(process)
        row[3].text = str(queuing)

    table.style = 'Colorful List'

    return doc

def doc_add_shrage_sol(N,a,d,q):
    '''
    add to the previous document, the solution obtained with schrage heurisitc
    '''
    document = creat_doc_data(N,a,d,q)
    
    document.add_page_break()
    document.add_heading('Solve with the Schrage Heuristic', 0)
    solution , fig = solve_schrage_heuristic(N,a,d,q , show_output = False ,graphics = False , larg = 10 , haut = 10)
    schedule = ["Job "+str(i) for i in solution['SCHD']]
    value = solution['UB']
    run_time = solution['time']

    # creation of the table of the solution
    document.add_heading('The raw data : ', 1)
    table = document.add_table(rows=1, cols=2)

    

    row = table.rows[0].cells

    row = table.add_row().cells
    row[0].text = 'Total time'
    row[1].text = str(value)

    row = table.add_row().cells
    row[0].text = "Schedule"
    row[1].text = str(schedule)

    row = table.add_row().cells
    row[0].text = "Running time"
    row[1].text = str(np.round(run_time,3))
    
    table.style = 'Colorful List'

    document.add_heading('The real Schedule', 1)
    fig = visu_calender(N,a,d,q,solution['SCHD'])

    plt.savefig('client_report/buff.png')
    document.add_picture('client_report/buff.png', width=Inches(7))

    return document

def doc_add_mip_sol(N,a,d,q):
    '''
    add to the previous document the solution with the mip
    '''
    document = doc_add_shrage_sol(N,a,d,q)
    
    document.add_page_break()
    document.add_heading('Solve with the Pulp Solver', 0)
    solution = solve_pulp_model(N,a,d,q , show_output = False)
    schedule = ["Job "+str(i) for i in solution['SCHD']]
    value = solution['OBJ']
    run_time = solution['RUN_TIME']

    # creation of the table of the solution
    document.add_heading('The raw data : ', 1)
    table = document.add_table(rows=1, cols=2)

    

    row = table.rows[0].cells

    row = table.add_row().cells
    row[0].text = 'Total time'
    row[1].text = str(value)

    row = table.add_row().cells
    row[0].text = "Schedule"
    row[1].text = str(schedule)

    row = table.add_row().cells
    row[0].text = "Running time"
    row[1].text = str(np.round(run_time,3))
    
    table.style = 'Colorful List'

    document.add_heading('The real Schedule', 1)
    fig = visu_calender(N,a,d,q,solution['SCHD'])

    plt.savefig('client_report/buff.png')
    document.add_picture('client_report/buff.png', width=Inches(7))

    return document



def create_save_doc(N,a,d,q , name = "default") :
    '''
    input : N, a, d, q --> data of the problem
            name       --> name of the document (without the .docx)
    '''
    name += ".docx" # --> docx format for the document
    doc = doc_add_mip_sol(N,a,d,q)
    doc.save("client_report/"+name)
    
